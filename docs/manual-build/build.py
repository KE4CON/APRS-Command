# -*- coding: utf-8 -*-
"""
Build docs/USER_MANUAL.docx from the chapter content in this folder.

Run:  python docs/manual-build/build.py
Output: docs/USER_MANUAL.docx  (styled; screenshot placeholders; auto TOC)

Chapter numbers are assigned by registry order (see CHAPTERS at the bottom) and
are PROVISIONAL until the full Table of Contents is locked.
"""
import os
import style as S
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===========================================================================
#  FRONT MATTER
# ===========================================================================
def title_page(doc):
    import datetime
    S.cover(
        doc,
        kicker="APRS COMMAND  ·  AMATEUR RADIO APRS",
        big_title="APRS COMMAND",
        subtitle="Situational Awareness for Amateur Radio",
        doc_kind="COMPLETE USER MANUAL",
        version="v1.0",
        tagline="APRS-IS & RF  ·  Receive-First  ·  Transmit-Safe  ·  EmComm-Ready",
        author="James Rospopo  ·  KE4CON",
        date_str=datetime.date.today().strftime("%B %d, %Y"),
    )


def philosophy(doc):
    S.page_break(doc)
    S.section_title(doc, "A Note on Philosophy")
    S.body(doc, "APRS Command exists because of two people and one promise.")
    S.body(doc, [("Bob Bruninga WB4APR", 'b'), " began building what became APRS in the early 1980s — his first "
        "position-mapping software ran on an Apple II in 1982, and a 1984 version tracked riders in a 100-mile endurance "
        "run — and he spent the following decades refining it into the system we use today. He was always clear about "
        "what it was for. APRS is not a tracking system. It is a ",
        ("situational awareness tool", 'bi'), " — a way for amateur radio operators to share real-time tactical "
        "information about what is happening in an area: weather, resources, objects, messages, coverage — a common "
        "operating picture. Bob wanted operators — whether at an emergency scene, a public service event, or a Field "
        "Day site — to look at a map and immediately understand the situation. That vision is the foundation of this program."])
    S.body(doc, [("Roger Barker G4IDE", 'b'), " built UI-View32, a widely used APRS client many operators relied on. "
        "When Roger passed in 2004, the source code was destroyed per his wishes, and UI-View32 could never be updated, "
        "fixed, or ported again. A program many operators depended on was frozen in time and slowly left behind by the "
        "operating systems it ran on — because it was closed source."])
    S.body(doc, [("APRS Command is released under GPL v3", 'b'), " — the GNU General Public License, version 3. The source "
        "code is always available; anyone can study it, fix it, improve it, and distribute it. No one can ever take it "
        "closed source. No one can ever destroy the source code. If the original author is gone tomorrow, any ham radio "
        "operator in the world can pick it up and carry it forward. That is the promise this license makes — to Bob's "
        "vision, to Roger's legacy, and to the amateur radio community that depends on these tools when it matters most."])
    S.body(doc, ["Use this program as Bob intended: not just to watch dots move on a map, but to understand what is "
        "happening in your operating area. Use it during emergency activations and exercises to build and share a common "
        "operating picture. Use it at public service events — hamfests, Field Day, parades, marathons, search and rescue "
        "— wherever operators need to coordinate. Use it to serve your community. That is what APRS was made for."])
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    qr = q.add_run("73 de KE4CON"); qr.italic = True; qr.font.size = Pt(11); qr.font.color.rgb = S.ACCENT


def how_to_use(doc):
    S.page_break(doc)
    S.section_title(doc, "How to Use This Manual")
    S.body(doc, "This manual is written to be read start-to-finish by a brand-new operator, or dipped into one chapter "
        "at a time. Every feature has its own chapter. Every chapter walks through the feature step by step, with nothing "
        "assumed. If you ever wonder what a button, panel, or field does, there is a chapter that explains it.")
    S.h2(doc, "The boxes you will see")
    S.callout(doc, "important", "IMPORTANT", "Something you must not overlook — usually about safety or transmitting. Read these carefully.")
    S.callout(doc, "note", "NOTE", "A helpful clarification or a detail about how something behaves.")
    S.callout(doc, "tip", "TIP", "A shortcut or a better way to do something once you are comfortable.")
    S.callout(doc, "warning", "WARNING", "An action that can have real-world consequences — for example, transmitting on the air.")
    S.h2(doc, "Screenshots")
    S.body(doc, ["Dashed boxes marked ", ("SCREENSHOT", 'b'), " show where a picture of the screen will appear in the "
        "finished manual. Each one describes exactly what it will show, so you can follow along even before the images are added."])
    S.h2(doc, "What things are called")
    S.body(doc, ["This manual uses one consistent name for each part of the screen — the same names the program itself "
        "uses. The main map screen is the ", ("map page", 'b'), ". The dark strip across the very top is the ",
        ("title bar", 'b'), ". The vertical strip of icon buttons down the left edge is the ", ("icon sidebar", 'b'),
        ". The strip along the very bottom is the ", ("status bar", 'b'), ". Any feature that opens in its own movable "
        "window is a ", ("floating panel", 'b'), ". These names are introduced with pictures in the map-page tour chapter."])
    S.callout(doc, "note", "A note on safety-first design.",
        "APRS Command starts up receive-only. It will not transmit anything until you deliberately set up your station "
        "and turn transmitting on. Throughout this manual, anything that can put a signal on the air is clearly flagged.")


def amendments_register(doc):
    S.page_break(doc)
    S.section_title(doc, "Amendments Register")
    S.body(doc, ["This manual uses ", ("stable chapter and section numbers", 'b'), ". When the program changes, we do not "
        "reprint the whole book — we publish a short, dated ", ("amendment", 'b'), " that revises or adds specific sections. "
        "Each amendment is listed in the table below. To keep a printed copy current, print the amendment and file it with "
        "the book; the register tells you what has changed and when."])
    S.body(doc, [("Tags: ", 'b'), ("AMENDS §X.Y", 'c'), " revises an existing section;  ", ("ADDS §Z", 'c'),
        " introduces a new one."])
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    widths = [Inches(1.1), Inches(1.0), Inches(1.4), Inches(3.0)]
    hdr = tbl.rows[0].cells
    for c, txt in zip(hdr, ["Date", "Tag", "Section", "Summary"]):
        c.width = widths[hdr.index(c)] if c in hdr else Inches(1)
        run = c.paragraphs[0].add_run(txt); run.bold = True; run.font.size = Pt(10); run.font.color.rgb = S.WHITE
    row = tbl.add_row().cells
    row[0].paragraphs[0].add_run("—")
    row[3].paragraphs[0].add_run("No amendments yet. This is the original edition.")
    for r in tbl.rows:
        for i, c in enumerate(r.cells):
            c.width = widths[i]
    doc.add_paragraph()
    S.body(doc, [("How this book is maintained. ", 'b'), "Chapter and section numbers, once assigned, are never reused or "
        "renumbered. New material is added at the end of the relevant part with the next free number, so a reference like "
        "“see §12.3” always points to the same place, in every edition. This is what makes small, printable "
        "amendments possible instead of a full reprint."])


def contents(doc):
    S.page_break(doc)
    h = doc.add_paragraph(); hr = h.add_run("Contents")
    hr.font.name = 'Segoe UI Semibold'; hr.font.size = Pt(24); hr.font.color.rgb = S.ACCENT
    rule = doc.add_paragraph(); S.par_border(rule, 'bottom', sz=12, color=S.ACCENT_HEX, space=8)
    note = doc.add_paragraph()
    nr = note.add_run("If the list below is blank or out of date, click it and press F9 (or right-click → Update Field) "
                      "to rebuild it.")
    nr.italic = True; nr.font.size = Pt(9.5); nr.font.color.rgb = S.MUTED
    S.toc(doc)


# ===========================================================================
#  CHAPTERS  (each: def chapter(doc, n))
# ===========================================================================
def ch_welcome(doc, n):
    S.chapter_open(doc, n, "Welcome to APRS Command",
        "What the program is, what it does, and the one rule that keeps you safe.",
        ["What APRS Command is", "What it can do", "What it will not do on its own",
         "The transmit-safety promise", "What you need to run it"])

    S.h1(doc, "What APRS Command Is")
    S.body(doc, "APRS Command is a desktop program for amateur-radio APRS operation. It runs on Windows, macOS, Linux, "
        "and the Raspberry Pi. It listens for APRS packets, shows the stations and weather it hears on a map, and gives "
        "you organized tools for messages, objects, alerts, logging, replay, and — when you choose to set it up — "
        "transmitting.")
    S.body(doc, ["APRS Command is built around one idea: ", ("safe, receive-first operation", 'b'), ". Out of the box it "
        "only listens. Nothing goes on the air until you deliberately configure your station and turn transmitting on."])

    S.h1(doc, "What APRS Command Can Do")
    S.bullets(doc, [
        "Receive APRS packets from the internet (APRS-IS) and from radio hardware (TCP KISS, Serial KISS, Direwolf, and AGWPE).",
        "Decode positions, messages, objects, items, weather, telemetry, status, and station-capability packets.",
        "Show every station it hears on a live map, and in a sortable station list.",
        "Show the raw packets as they arrive, so you can confirm what is really being received.",
        "Display weather from APRS weather stations, plus a national-weather radar overlay and alerts.",
        "Organize private messages, bulletins, announcements, and queries.",
        "Create and preview objects and items.",
        "Watch for conditions you care about with alerts and geofences.",
        "Replay recorded traffic, run a built-in simulation, and practice in a training mode — none of which touch the air.",
        "Prepare offline map tiles for use where there is no internet.",
    ])

    S.h1(doc, "What APRS Command Does Not Do On Its Own")
    S.body(doc, "By default, APRS Command does not transmit anything. It will not put a signal on the air, connect to "
        "APRS-IS as a sender, beacon your position, act as an iGate or digipeater, send messages or objects, or beacon "
        "weather — until you set those things up on purpose.")
    S.callout(doc, "important", "IMPORTANT — receive first.",
        "Everything that could transmit is turned off when you first install the program. This is deliberate. Get "
        "comfortable listening and watching the map before you even think about transmitting.")

    S.h1(doc, "The Transmit-Safety Promise")
    S.body(doc, ["Two indicators on the ", ("title bar", 'b'), " (the dark strip across the very top of the window) always "
        "tell you your transmit state. These are the ", ("status badges", 'b'), ":"])
    S.bullets(doc, [
        [("The APRS-IS badge", 'b'), " — whether you are connected to the APRS internet network, and whether that "
         "connection is receive-only or able to send."],
        [("The TX badge", 'b'), " — your on-the-air transmit state. When it reads ", ("TX Disabled", 'c'),
         ", the program cannot key a radio, no matter what else is going on."],
    ])
    S.body(doc, "You will get to know these two badges well. Whenever you are unsure whether the program can transmit, "
        "glance at the title bar — the badges never lie, and a whole chapter later in this manual is devoted to them.")
    S.callout(doc, "warning", "WARNING — you are the licensed operator.",
        "When you do enable transmitting, you are responsible for legal amateur-radio operation: your callsign, SSID, "
        "path, beacon interval, and local rules are yours to get right. This manual will walk you through each of those "
        "before you ever key up. Always test receive-only first.")

    S.h1(doc, "System Requirements")
    S.body(doc, "APRS Command is a 64-bit desktop application. Here is exactly what it runs on — and what it will not.")

    S.h2(doc, "Supported operating systems")
    S.table(doc,
        ["Platform", "Minimum version", "Runs on"],
        [
            ["Windows", "Windows 10 or Windows 11 (64-bit)", "64-bit PCs and laptops"],
            ["macOS", "macOS 14 (Sonoma) or later", "Apple Silicon (M1–M4) and Intel Macs"],
            ["Linux", "Ubuntu 22.04+, Debian 12+, Fedora, or RHEL 8+ (64-bit)", "x64 or ARM64, with a desktop"],
            ["Raspberry Pi", "Raspberry Pi OS 64-bit (Bookworm)", "Pi 3, 4, 5, 400, Zero 2 W — Pi 4 / 5 recommended"],
        ],
        [Inches(1.15), Inches(2.75), Inches(2.6)])
    S.callout(doc, "note", "Where these numbers come from.",
        "These minimums are set by the .NET 10 runtime that APRS Command is built on. On Linux you also need a graphical "
        "desktop (X11 or Wayland) and glibc 2.27 or newer — any distribution from 2022 onward meets this.")

    S.h2(doc, "Memory (RAM)")
    S.body(doc, [("2 GB minimum; 4 GB or more recommended.", 'b'), " The live map is the most memory-hungry part of the "
        "program. On a desktop or laptop this is never a concern. On a Raspberry Pi, aim for 2 GB or more — the 1 GB Pi 3 "
        "and the 512 MB Pi Zero 2 W will run APRS Command, but the map will feel slow."])

    S.h2(doc, "Storage")
    S.body(doc, ["About ", ("300 MB", 'b'), " for the program itself. Leave extra room for your logs and — if you download "
        "offline map areas — the map cache, which can grow to hundreds of megabytes or more depending on how large an area "
        "and how much detail you save."])

    S.h2(doc, "Display")
    S.body(doc, ["A normal desktop screen. The map-first layout is comfortable at ", ("1280 × 768 or larger", 'b'),
        ". APRS Command is a windowed desktop program — it does not run “headless” (with no screen), and it is "
        "not a phone or tablet app."])

    S.h2(doc, "Optional — for RF and live data")
    S.bullets(doc, [
        [("Internet access", 'b'), " for APRS-IS receive and online map tiles. You can also run fully offline with "
         "downloaded map areas."],
        [("Radio hardware", 'b'), " — a TNC, a Direwolf or AGWPE sound-card modem, or a serial/TCP link — to receive "
         "(and later transmit) over the air instead of the internet."],
        [("No administrator account", 'b'), " is needed to run APRS Command; only the installer needs it, to copy the "
         "program into place."],
    ])

    S.h2(doc, "What APRS Command will not run on")
    S.bullets(doc, [
        [("32-bit operating systems", 'b'), " — 32-bit Windows, or the 32-bit version of Raspberry Pi OS or Linux. "
         "APRS Command ships as 64-bit only."],
        [("Older ARM boards", 'b'), " — the original Raspberry Pi, the Pi 1, the first Pi Zero and Zero W, and the early "
         "Pi 2 (v1.1). Their processors are 32-bit and too old for the runtime."],
        [("Operating systems older than the minimums above", 'b'), " — Windows 8.1 or earlier, macOS 13 or earlier, or "
         "Linux distributions from before 2022."],
        [("Phones and tablets", 'b'), " (iOS, Android) and headless servers with no graphical desktop."],
    ])
    S.callout(doc, "tip", "Not sure about your Pi?",
        "If it can install and boot the 64-bit Raspberry Pi OS, it can run APRS Command — that means the Pi 3 and newer, "
        "and the Pi Zero 2 W. If only the 32-bit OS will install, the board is too old.")

    S.body(doc, ["With the requirements met, the next chapters get you from a fresh install to a working map: ",
        ("Installing APRS Command", 'i'), ", then ", ("First Launch & First-Run Setup", 'i'), ", then a guided ",
        ("tour of the map page", 'i'), "."])


def ch_installing(doc, n):
    S.chapter_open(doc, n, "Installing APRS Command",
        "Getting the program onto your computer — Windows, macOS, Linux, or a Raspberry Pi.",
        ["Two ways to install", "Windows", "macOS", "Linux", "Raspberry Pi",
         "Signed vs. unsigned: what you'll see", "Running from source"])

    S.callout(doc, "note", "Signed or unsigned — this chapter covers both.",
        "Depending on the build you download, APRS Command may be digitally signed (it installs cleanly and shows the "
        "author's name) or unsigned (your computer shows a one-time security caution you safely click past). It is the same "
        "program either way. Each platform's steps below tell you what to expect in both cases, and "
        "“Signed vs. Unsigned: What You'll See” near the end explains the difference in plain terms.")

    S.h1(doc, "Before You Start")
    S.body(doc, ["Make sure your computer meets the ", ("System Requirements", 'i'), " in the previous chapter — in short, "
        "a 64-bit Windows, macOS 14+, modern Linux, or a 64-bit Raspberry Pi. You do ", ("not", 'i'), " need an "
        "administrator account to run APRS Command; only the installer needs one, to copy the program into place."])
    S.callout(doc, "note", "Two ways to install.",
        "Every platform offers an installer (the simplest choice — double-click and go) and a portable archive (a .zip or "
        ".tar.gz you unpack and run in place, with no installation). Pick whichever you prefer; both give you the same program.")

    S.h1(doc, "Windows")
    S.h2(doc, "Installer (recommended)")
    S.steps(doc, [
        ["Download the Windows installer — ", ("APRSCommand-…-windows-x64-Setup.exe", 'c'), "."],
        ["Double-click it to run."],
        ["The first time you run it, watch for one of two things — both are fine. If this build is ", ("signed", 'b'),
         ", a ", ("User Account Control", 'b'), " box may appear listing ", ("Verified publisher: James Rospopo (KE4CON)", 'b'),
         " — click ", ("Yes", 'b'), ". If instead you see ", ("“Windows protected your PC,”", 'b'),
         " this copy is unsigned — click ", ("More info", 'b'), ", then ", ("Run anyway", 'b'),
         " (safe; see “Signed vs. Unsigned” below)."],
        ["APRS Command installs to ", ("C:\\Program Files\\APRS Command", 'c'), " with Start-menu and desktop shortcuts."],
    ])
    S.body(doc, ["To remove it later: ", ("Settings → Apps → APRS Command → Uninstall", 'b'), "."])
    S.h2(doc, "Portable (no installation)")
    S.steps(doc, [
        ["Download ", ("APRSCommand-…-windows-x64.zip", 'c'), " and extract it to any folder."],
        ["Run ", ("Aprs.Desktop.exe", 'c'), ". If a ", ("“Windows protected your PC”", 'b'),
         " prompt appears (that means an unsigned copy), click ", ("More info → Run anyway", 'b'),
         ". A signed copy just runs."],
    ])
    S.screenshot(doc, "Side by side: a signed User Account Control prompt showing the verified publisher, and the unsigned “Windows protected your PC” prompt with More info / Run anyway")

    S.h1(doc, "macOS")
    S.h2(doc, "Installer (.dmg, recommended)")
    S.steps(doc, [
        ["Download the disk image for your Mac — ", ("…-macos-arm64.dmg", 'c'), " for Apple Silicon (M1–M4) or ",
         ("…-macos-x64.dmg", 'c'), " for an Intel Mac."],
        ["Double-click the ", (".dmg", 'c'), " to open it, then drag ", ("APRS Command", 'b'), " into your ",
         ("Applications", 'b'), " folder."],
        ["The first time you open it, watch for one of two things. If this build is ", ("signed and notarized", 'b'),
         ", just ", ("double-click", 'b'), " it — it opens normally (you may get a one-time “downloaded from the Internet — "
         "are you sure you want to open it?” with an ", ("Open", 'b'), " button; click ", ("Open", 'b'),
         "). If instead macOS says it ", ("“cannot be opened because Apple cannot check it,”", 'b'),
         " this copy is unsigned — ", ("right-click", 'b'), " (or Control-click) the app, choose ", ("Open", 'b'),
         ", then ", ("Open", 'b'), " again. macOS remembers it, and it launches normally afterward."],
    ])
    S.callout(doc, "note", "The right-click step is only for an unsigned copy.",
        "A signed, notarized build opens with a normal double-click — no dance needed. The right-click → Open step is the "
        "standard one-time approval only when the copy is unsigned. See “Signed vs. Unsigned: What You'll See” below.")
    S.h2(doc, "Portable (.zip)")
    S.body(doc, ["Unzip the archive, then in Terminal run ", ("xattr -cr .", 'c'), " inside the extracted folder (to clear "
        "the download-quarantine flag) and launch ", ("./Aprs.Desktop", 'c'), "."])

    S.h1(doc, "Linux")
    S.body(doc, "APRS Command comes in three Linux formats. First, pick the download that matches your computer:")
    S.bullets(doc, [
        [("amd64", 'c'), " — a standard 64-bit PC or laptop (Intel or AMD processor)."],
        [("arm64", 'c'), " — a Raspberry Pi or other ARM-based computer."],
    ])
    S.callout(doc, "note", "About the version number in these commands.",
        "The commands below show an example version, 1.0.0. Type the exact name of the file you actually downloaded — "
        "your version number may be different. (In a terminal, you can type the first few letters and press Tab to "
        "complete the file name for you.)")

    S.h2(doc, "Debian, Ubuntu, Mint, or Raspberry Pi OS  (.deb)")
    S.body(doc, "Open a terminal in the folder where the file was downloaded, then run the command for your computer.")
    S.body(doc, [("On a PC or laptop", 'b'), " (amd64):"])
    S.code_block(doc, "sudo dpkg -i aprs-command_1.0.0_amd64.deb")
    S.body(doc, [("On a Raspberry Pi or ARM computer", 'b'), " (arm64):"])
    S.code_block(doc, "sudo dpkg -i aprs-command_1.0.0_arm64.deb")

    S.h2(doc, "Fedora, RHEL, or openSUSE  (.rpm)")
    S.body(doc, "Open a terminal in the download folder, then run the command for your computer.")
    S.body(doc, [("On a PC or laptop", 'b'), " (x86_64):"])
    S.code_block(doc, "sudo rpm -i aprs-command-1.0.0-1.x86_64.rpm")
    S.body(doc, [("On a Raspberry Pi or ARM computer", 'b'), " (aarch64):"])
    S.code_block(doc, "sudo rpm -i aprs-command-1.0.0-1.aarch64.rpm")

    S.h2(doc, "Any distribution — portable archive  (.tar.gz)")
    S.body(doc, "No installation needed. Unpack the archive and run the program in place:")
    S.code_block(doc, [
        "tar -xzf APRSCommand-1.0.0-linux-x64.tar.gz",
        "cd APRS-Command-linux-x64",
        "chmod +x Aprs.Desktop",
        "./Aprs.Desktop",
    ])

    S.body(doc, ["An installed package (.deb or .rpm) places the program at ", ("/opt/aprs-command/", 'c'), ", adds an ",
        ("aprs-command", 'c'), " command you can run from a terminal, and creates an entry in your application menu."])
    S.callout(doc, "tip", "Using a hardware TNC over serial?",
        "On Linux and macOS your user must belong to the “dialout” group to open serial ports: run "
        "“sudo usermod -aG dialout $USER”, then log out and back in. On Windows no extra step is needed. "
        "This is covered fully in the RF / TNC Connections chapter.")

    S.h1(doc, "Raspberry Pi")
    S.steps(doc, [
        ["Install the ", ("64-bit", 'b'), " Raspberry Pi OS (the tested version is “Bookworm”). The 32-bit OS will not run APRS Command."],
        ["Use the ARM64 package: ", ("sudo dpkg -i aprs-command_…_arm64.deb", 'c'), "."],
        ["Launch it from the application menu, or by running ", ("aprs-command", 'c'), "."],
    ])
    S.callout(doc, "note", "NOTE", "Keep your logs and the offline map cache on storage with room to spare — on a Pi, "
        "avoid filling the boot SD card with downloaded map tiles. A Pi 4 or Pi 5 gives the smoothest map.")

    S.h1(doc, "Signed vs. Unsigned: What You'll See")
    S.body(doc, "“Code signing” is a digital seal that tells your computer who made a program and proves it has not been "
        "altered since. When a program is signed, Windows and macOS trust it and show you the author's name. When it is "
        "not signed, they show a one-time caution the first time you run it — not because anything is wrong, but because "
        "they cannot confirm the author automatically.")
    S.body(doc, ["APRS Command is a free, open-source amateur-radio project. Depending on the download, a build may be ",
        ("signed", 'b'), " — published under the author's verified name, ", ("James Rospopo, KE4CON", 'b'),
        " — or ", ("unsigned", 'b'), ". Both are the identical program, from the same public source code. The only "
        "difference is whether your computer greets it by name or with a caution. Here is exactly what to expect, and what "
        "to do, either way:"])
    S.table(doc,
        ["Your computer", "If the copy is signed", "If the copy is unsigned — and what to do"],
        [
            ["Windows",
             "Installs normally. A “User Account Control” box may show “Verified publisher: James Rospopo (KE4CON)” — click Yes.",
             "A “Windows protected your PC” box appears. Click More info, then Run anyway. Safe — it only means this copy is not signed."],
            ["macOS",
             "Opens with a normal double-click. You may get a one-time “downloaded from the Internet — open?” — click Open.",
             "macOS says it “cannot be opened because Apple cannot check it.” Right-click (or Control-click) the app → Open → Open."],
            ["Linux / Raspberry Pi",
             "No warning — Linux does not gate apps this way. Install and run normally.",
             "No warning either; install and run normally."],
        ],
        [Inches(1.2), Inches(2.65), Inches(2.65)])
    S.callout(doc, "tip", "Want to confirm you have a genuine copy?",
        "Always download from the official releases page: github.com/KE4CON/APRS-Command/releases. On Windows you can also "
        "right-click the installer → Properties → the Digital Signatures tab to see who signed it; on macOS, a signed build "
        "simply opens without the “cannot be checked” message.")
    S.callout(doc, "note", "Both are safe — and this can change over time.",
        "Whether a given release is signed may vary (signing certificates renew periodically, and this is an independent "
        "project). If a build you expected to be signed shows a caution, or the reverse, nothing is wrong — just follow the "
        "matching column above. Either way it is the same program from the same source.")

    S.h1(doc, "Running from Source (Optional)")
    S.body(doc, "If you would rather build it yourself — for example to try the very latest changes — install the "
        ".NET 10 SDK, then run:")
    S.code_block(doc, [
        "git clone https://github.com/KE4CON/APRS-Command.git",
        "cd APRS-Command",
        "dotnet run --project src/Aprs.Desktop/Aprs.Desktop.csproj",
    ])
    S.body(doc, ["With APRS Command installed, the next chapter walks through your ", ("first launch and first-run setup", 'i'), "."])


def ch_maptour(doc, n):
    S.chapter_open(doc, n, "A Tour of the Map Page",
        "Every part of the main screen, named — so the rest of this manual always makes sense.",
        ["The map page at a glance", "Title bar & status badges", "The menu bar",
         "The icon sidebar", "The map area", "Panels that appear when needed",
         "The status bar", "How other features open"])

    S.h1(doc, "The Map Page at a Glance")
    S.body(doc, ["When APRS Command opens, you land on the ", ("map page", 'b'), " — the main screen you will spend most of "
        "your time on. It has a small number of fixed parts, and this chapter names every one. These names are used "
        "consistently throughout the manual, so a few minutes here pays off in every later chapter."])
    S.screenshot(doc, "The whole map page with each region labeled: title bar, menu bar, icon sidebar, map area, status bar")

    S.h1(doc, "The Title Bar & Status Badges")
    S.body(doc, ["The ", ("title bar", 'b'), " is the dark strip across the very top. It shows the app name, ",
        ("APRS Command", 'b'), ", and the subtitle “APRS station map.” At its right end sit two ",
        ("status badges", 'b'), " that are always visible:"])
    S.bullets(doc, [
        [("The APRS-IS badge", 'b'), " — your connection to the APRS internet network: ", ("APRS-IS Offline", 'c'),
         " or connected."],
        [("The TX badge", 'b'), " — your on-the-air transmit state. ", ("TX Disabled", 'c'), " means the program cannot "
         "key a radio, no matter what else is happening."],
    ])
    S.callout(doc, "important", "Glance here whenever you are unsure.",
        "These two badges are your at-a-glance truth about whether APRS Command can transmit. They are always in the same "
        "place, and they never lie. A whole later chapter is devoted to transmit safety and the TX badge.")

    S.h1(doc, "The Menu Bar")
    S.body(doc, ["Just below the title bar is the ", ("menu bar", 'b'), ". Almost every window and feature in APRS Command "
        "opens from one of its seven menus:"])
    S.table(doc,
        ["Menu", "What you'll find there"],
        [
            ["Settings", "Opens the settings window (station identity, connections, first-run setup, safety, and more)."],
            ["View", "Station List, Raw Packets, Telemetry, Events, the Packet Statistics dashboard, and the dark-mode toggle."],
            ["Map", "The drawing tools, weather (stations, alerts, radar), offline map download, coverage prediction, elevation profile, and the frequency reference."],
            ["Messages", "The Message Center, message broadcast, scheduled messages, and the receipts dashboard."],
            ["Operate", "Net Control, the net script editor, objects, session templates, scheduled and shadow beacons, and alerts."],
            ["Tools", "Event Bus, Replay, RF Diagnostics, the After-Action report, the Mobile Companion web view, and Exercise Mode."],
            ["Help", "The in-app Help viewer, the keyboard-shortcut list, and the About window."],
        ],
        [Inches(1.0), Inches(5.5)])
    S.body(doc, "Each of those features has its own chapter later in this manual; here we are only naming where they live.")

    S.h1(doc, "The Icon Sidebar")
    S.body(doc, ["Down the left edge is the ", ("icon sidebar", 'b'), " — a vertical strip of icon buttons for the tools "
        "you reach for most. The buttons have no printed labels; ", ("hover over any icon to see its name in a tooltip", 'b'),
        ". A divider separates the map tools (top) from the operator tools (below)."])
    S.table(doc,
        ["Icon", "Name", "What it does"],
        [
            ["⌂", "Home", "Reset the map to the default overview."],
            ["◎", "Center on my station", "Recenter the map on your own station."],
            ["🔍", "Find station", "Search for a callsign and jump to it."],
            ["📏", "Measure distance", "Measure the distance between points on the map."],
            ["🗺", "Map layer", "Switch the base-map layer (see below)."],
            ["📡", "Beacon Now", [("Transmit your position immediately. ", ""), ("(Only works once you have set up and enabled transmit.)", 'i')]],
            ["🔔", "Alerts", "Show your alert status."],
            ["⊙", "Range rings", "Toggle distance range rings around a point."],
            ["🛤", "Trails", "Toggle station movement trails."],
            ["🌧", "Radar", "Toggle the weather-radar overlay."],
        ],
        [Inches(0.5), Inches(1.9), Inches(4.1)])
    S.callout(doc, "warning", "WARNING — Beacon Now transmits.",
        "The Beacon Now button (📡) puts your position on the air the moment transmit is enabled. While you are learning "
        "receive-first, it does nothing — but treat it with respect once your station is configured.")

    S.h1(doc, "The Map Area")
    S.body(doc, ["The rest of the window is the ", ("map area", 'b'), " — the live map where stations, objects, and weather "
        "appear. In its top-left corner is the ", ("base-map selector", 'b'), ", a dropdown that switches the background map:"])
    S.bullets(doc, [
        [("OpenStreetMap", 'b'), " — a clear street map (the default)."],
        [("USGS Topo", 'b'), " — topographic contour maps."],
        [("USGS Imagery", 'b'), " — aerial/satellite photography."],
        [("USGS Imagery + Topo", 'b'), " — aerial imagery with topographic labels on top."],
    ])
    S.body(doc, "Click a station, object, or weather marker to select it. Moving around the map (panning and zooming) and "
        "each base map are covered in the next chapter.")

    S.h1(doc, "Panels That Appear Only When Needed")
    S.body(doc, "Several small panels appear on the map only when they are relevant, then disappear again:")
    S.bullets(doc, [
        [("Station-details panel", 'b'), " (bottom-left) — the selected station's details, shown when you click a marker."],
        [("Radar scrubber", 'b'), " (bottom-center) — step or play through radar frames, shown when weather radar is on."],
        [("Object-placement overlay", 'b'), " (top-right) — guides you while placing or moving an object."],
        [("Draw-mode banner", 'b'), " (across the top) — shown while a drawing tool is active (see the Drawing chapter)."],
        [("Weather-alert banner", 'b'), " (full width) — announces an active National Weather Service alert; click it for details."],
    ])

    S.h1(doc, "The Status Bar")
    S.body(doc, ["Along the very bottom is the ", ("status bar", 'b'), ", a quiet summary of the program's state — for "
        "example ", ("Ready", 'c'), ", ", ("APRS-IS Disconnected", 'c'), " or ", ("Connected", 'c'), ", and ",
        ("RF TX Disabled", 'c'), " or ", ("Enabled", 'c'), ". It is a second, always-present confirmation of your transmit and connection state."])

    S.h1(doc, "How the Other Features Open")
    S.body(doc, ["Everything beyond the map itself — the Station List, Messages, Objects, Weather, Replay, and the rest — "
        "opens from the menus as its own ", ("floating panel", 'b'), ": a movable window you can drag by its ",
        ("panel title bar", 'b'), " (the colored strip with the grip dots) and close with its ", ("✕", 'b'),
        ". You can have several open at once and arrange them around the map however suits you."])
    S.callout(doc, "tip", "Prefer a dark screen?",
        "Choose View → Toggle Dark Mode to switch between light and dark themes at any time — handy for night operating.")
    S.body(doc, ["Now that the map page has names for all its parts, the next chapters put them to work — starting with ",
        ("moving around the map and choosing base maps", 'i'), "."])


def ch_drawing(doc, n):
    S.chapter_open(doc, n, "Drawing on the Map",
        "Annotate the map with lines, shapes, circles, and text — with color, fill, width, and real-world measurements.",
        ["What the drawing tools are for", "Opening the drawing toolbar", "A tour of the toolbar",
         "Choosing color, fill, and width", "Lines, polygons, and circles",
         "Adding, sizing, and moving text", "The eyedropper", "Measurements",
         "Erasing, clearing, and exiting", "Importing & exporting (GPX / KML)", "Troubleshooting"])

    S.h1(doc, "What the Drawing Tools Are For")
    S.body(doc, "The drawing tools let you mark directly on the map — lines, filled shapes, and circles — to turn a plain "
        "map into a planning and situational-awareness picture. Operators use them to outline a search sector, mark a "
        "coverage area or a net boundary, sketch a route, or flag a staging area during an event.")
    S.callout(doc, "important", "IMPORTANT — drawings stay on your screen.",
        "Everything you draw is a local annotation on your own map only. It is never transmitted over the air (RF) or to "
        "the internet (APRS-IS), and no other station ever sees it. Draw freely — nothing you sketch goes anywhere.")

    S.h1(doc, "Opening the Drawing Toolbar")
    S.body(doc, ["All of the drawing is done from an on-map ", ("drawing toolbar", 'b'), " that appears only while you are "
        "drawing. To open it, choose ", ("Map → Draw → Drawing Toolbar", 'b'), " from the menu bar. The toolbar appears "
        "across the top of the map, and the map enters ", ("draw mode", 'b'), "."])
    S.body(doc, ["That same menu also holds ", ("Import GPX / KML", 'b'), " and ", ("Export Drawings", 'b'),
        " (covered later) — everything else lives on the toolbar."])
    S.screenshot(doc, "The Map → Draw menu open, showing Drawing Toolbar, Import, and Export")
    S.callout(doc, "note", "The map holds still while you draw.",
        "In draw mode, clicking and dragging draws on the map instead of moving it — so your points land exactly where you "
        "click and the map never jumps out from under you. You can still zoom with the mouse wheel. To pan the map again, "
        "click ✕ Exit on the toolbar to leave draw mode.")

    S.h1(doc, "A Tour of the Drawing Toolbar")
    S.body(doc, "The toolbar puts every drawing control in one place. The tool you are using is highlighted.")
    S.table(doc,
        ["Control", "What it does"],
        [
            ["Tool buttons", "Line, Polygon, Circle, Text, Resize/Move text, Eyedropper, and Erase. Click one to switch tools; the active tool lights up. (Switching tools never deletes what you have already drawn.)"],
            ["Color swatch", "Shows the current tool's color; click it to open the color wheel and choose any color."],
            ["Fill", "The fill for polygons and circles: None (outline only), Solid, or a hatch / dot pattern."],
            ["W (width)", "Line / outline thickness, from 1 to 12."],
            ["Measure", "Turns the length / diameter / area labels on shapes on or off."],
            ["Units", "Imperial or Metric, plus a 'Keep small units' toggle (stay in ft/acres · m/ha even for large shapes)."],
            ["Clear", "Deletes every drawing on the map."],
            ["✕ Exit", "Leaves draw mode (your drawings stay)."],
        ],
        [Inches(1.35), Inches(5.15)])
    S.screenshot(doc, "The drawing toolbar across the top of the map, with each control labeled")
    S.callout(doc, "important", "Each tool remembers its own style.",
        "Color, fill, and width are kept per tool. Your line can be a red 3-px line while your circle is a blue "
        "cross-hatch — switch tools and the toolbar shows that tool's settings. Changing a setting affects only the tool "
        "you are on.")

    S.h1(doc, "Choosing Color, Fill, and Width")
    S.body(doc, [("Color.", 'b'), " Click the color swatch on the toolbar to open the color wheel. Pick from the round "
        "spectrum, the Palette of preset swatches, or by typing a hex value — then click ", ("Done", 'b'),
        " (or click outside the wheel) to close it. The color applies to whatever you draw next with the current tool."])
    S.body(doc, [("Fill.", 'b'), " Polygons and circles can be filled. Click ", ("Fill", 'b'), " and choose ",
        ("None", 'b'), " (outline only), ", ("Solid", 'b'), " (a soft tint of the shape's color), or a pattern — "
        "Diagonal hatch, Cross-hatch, Horizontal, Vertical, or Dotted (handy for telling zones apart)."])
    S.body(doc, [("Width.", 'b'), " The ", ("W", 'b'), " box sets the line or outline thickness (1–12)."])
    S.screenshot(doc, "The color wheel open from the toolbar swatch, with its Done button")

    S.h1(doc, "Drawing a Line")
    S.steps(doc, [
        ["Click the ", ("Line", 'b'), " tool on the toolbar."],
        ["Click once on the map where the line should ", ("start", 'b'), "."],
        ["Click at each point the line should pass through — a preview segment follows your cursor."],
        ["When the line is complete, ", ("double-click", 'b'), " to finish it."],
    ])
    S.callout(doc, "note", "NOTE", "A line needs at least two points; a one-point line is discarded.")
    S.screenshot(doc, "A multi-point line being drawn, with the preview segment stretching to the cursor")

    S.h1(doc, "Drawing a Polygon")
    S.steps(doc, [
        ["Click the ", ("Polygon", 'b'), " tool."],
        ["Click each ", ("corner", 'b'), " of the area in turn — a preview follows your cursor."],
        ["Double-click to close the shape; the last corner joins back to the first."],
    ])
    S.callout(doc, "note", "NOTE", "A polygon needs at least three corners.")
    S.screenshot(doc, "A polygon being drawn around an area, with a hatch fill")

    S.h1(doc, "Drawing a Circle")
    S.steps(doc, [
        ["Click the ", ("Circle", 'b'), " tool."],
        ["Press and hold at the ", ("center", 'b'), " of the circle."],
        ["Drag ", ("outward", 'b'), " — a preview circle grows — and release to set the radius."],
    ])
    S.callout(doc, "tip", "TIP", "Drag out far enough to give the circle a real size. A quick click with no drag makes no circle.")
    S.screenshot(doc, "A circle being dragged outward from its center")

    S.h1(doc, "Adding Text")
    S.body(doc, "Text labels are sized by dragging, and they are anchored to the map — they grow and shrink with it as you "
        "zoom, so a label stays the right size for the area it marks.")
    S.steps(doc, [
        ["Click the ", ("Text", 'b'), " tool."],
        ["Press and ", ("drag", 'b'), " on the map — the further you drag, the larger the text."],
        ["Release, type your label, choose a ", ("Background", 'b'), " (None, White, Black, Yellow, or Light gray), and "
         "click ", ("Place", 'b'), "."],
    ])
    S.callout(doc, "note", "Text scales with the map.",
        "Unlike a fixed on-screen label, map text is tied to the ground: zoom in and it grows with the map; zoom out and it "
        "shrinks. That keeps it the right size for the area it labels at every zoom level.")
    S.screenshot(doc, "Placing text — dragging to size it, then the label dialog with the Background choice")

    S.h1(doc, "Resizing or Moving Text")
    S.steps(doc, [
        ["Click the ", ("Resize / Move text", 'b'), " tool on the toolbar."],
        ["Click a label to ", ("select", 'b'), " it — a small square handle appears."],
        ["Drag the ", ("handle", 'b'), " to resize the text, or drag the ", ("label itself", 'b'), " to move it."],
    ])
    S.screenshot(doc, "A selected text label showing its resize handle")

    S.h1(doc, "The Eyedropper")
    S.body(doc, "The eyedropper matches a color that is already on the map — a station symbol, the terrain, another drawing.")
    S.steps(doc, [
        ["Click the ", ("Eyedropper", 'b'), " tool."],
        ["Click anywhere on the map. The color under your click becomes the current tool's color, and you return to "
         "your drawing tool automatically."],
    ])

    S.h1(doc, "Measurements")
    S.body(doc, ["Turn on ", ("Measure", 'b'), " on the toolbar to label each shape with its real-world size:"])
    S.bullets(doc, [
        [("Line", 'b'), " — its length."],
        [("Circle", 'b'), " — its diameter and area."],
        [("Polygon", 'b'), " — its area."],
    ])
    S.body(doc, ["Use ", ("Units", 'b'), " to switch between ", ("Imperial", 'b'), " (feet, miles, acres) and ",
        ("Metric", 'b'), " (meters, kilometers, hectares) — the unit auto-scales to the size. The value is the true "
        "distance on the ground, stays correct as you zoom, is the same on every base map, and updates live while you draw."])
    S.body(doc, ["The Units menu also has a ", ("Keep small units", 'b'), " toggle: turn it on to stay in feet and acres "
        "(or meters and hectares) even for large shapes, instead of auto-scaling up to miles and square miles. The label "
        "appears just below each shape, clear of the outline."])
    S.callout(doc, "note", "NOTE", "Measurements are corrected for the map projection, so a mile measured near the poles "
        "reads the same as a mile at the equator — the on-screen size differs, but the real-world number is right.")
    S.screenshot(doc, "Shapes with measurement labels — a line's length, a circle's diameter and area, a polygon's area")

    S.h1(doc, "Erasing a Shape")
    S.steps(doc, [
        ["Click the ", ("Erase", 'b'), " tool."],
        ["Click on (or near) the shape to remove. It disappears; everything else stays."],
    ])
    S.body(doc, "Erase stays active, so you can remove several shapes one after another.")

    S.h1(doc, "Clearing Everything, and Leaving Draw Mode")
    S.body(doc, ["The toolbar's ", ("Clear", 'b'), " button deletes every drawing at once. The ", ("✕ Exit", 'b'),
        " button leaves draw mode and returns the map to normal panning — your drawings stay."])
    S.callout(doc, "important", "Clear vs. Exit.",
        "Clear deletes your drawings. ✕ Exit keeps them and just leaves draw mode. They sit next to each other on the "
        "toolbar, so click the one you mean. Clear cannot be undone — export first if you might want them later.")

    S.h1(doc, "Sharing Drawings with Other Apps: GPX & KML Files")
    S.body(doc, "APRS Command can read and write two standard map-data file formats, so the shapes on your map can come "
        "from — or go to — other mapping programs. Both are plain-text files that store geographic points, lines, and "
        "areas. You do not need to understand their insides; you just need to know which is which.")
    S.table(doc,
        ["Format", "Where it comes from", "What it is best at"],
        [
            ["GPX (.gpx)", "The GPS world — Garmin units, phone GPS apps, fitness watches, geocaching and hiking apps. "
             "GPX is the common language of GPS receivers.",
             "Waypoints (single marked points), tracks (a breadcrumb trail of where you actually went), and routes "
             "(a planned sequence of points)."],
            ["KML (.kml)", "The mapping/visualization world — it is the native format of Google Earth and Google Maps, "
             "and is common in ICS and public-safety mapping.",
             "Styled shapes and filled areas (a coverage zone, a search grid, a flood boundary) with colors, labels, "
             "and pop-up descriptions — built for displaying an annotated map."],
        ],
        [Inches(1.25), Inches(2.75), Inches(2.5)])
    S.body(doc, ["A quick rule of thumb: ", ("GPX = tracks and points from a GPS", 'b'), "; ",
        ("KML = annotated shapes for a map viewer", 'b'),
        ". Most tools read and write both, so for simple points and lines they are interchangeable."])
    S.callout(doc, "tip", "Real uses on the air.",
        "A neighboring EOC sends you a KML of the assigned search area — import it and the boundary appears right on your "
        "APRS map. Draw a coverage polygon yourself and export it as KML to open in Google Earth or hand to another team. "
        "Load a GPX full of waypoints (shelter sites, repeater locations, a planned drive route) and they show up as "
        "labeled points you can work against.")

    S.h2(doc, "Importing Shapes from a File")
    S.steps(doc, [
        ["Choose ", ("Map → Draw → Import GPX / KML…", 'b'), "."],
        ["Pick a ", (".gpx", 'b'), " or ", (".kml", 'b'), " file and open it. (The format is detected from the file "
         "itself, so a mis-named file still opens correctly.)"],
        ["The shapes appear on the map, added to whatever you have already drawn:"],
    ])
    S.bullets(doc, [
        [("From GPX:", 'b'), " waypoints become labeled points; tracks become blue lines; routes become gold lines."],
        [("From KML:", 'b'), " points become labeled points; lines become blue lines; filled areas (polygons) come in "
         "as their outer outline."],
    ])
    S.callout(doc, "note", "About imported colors and KMZ files.",
        "Imported shapes use APRS Command's own default colors (blue for lines/tracks, gold for routes/areas); any colors "
        "or fill styling saved in the original file are not carried in. And a .kmz file (a zipped .kml, the form Google "
        "Earth often saves) is not read directly — rename it to .zip, extract it, and open the .kml file inside.")

    S.h2(doc, "Exporting Your Drawings")
    S.steps(doc, [
        ["Choose ", ("Map → Draw → Export Drawings…", 'b'), "."],
        ["Pick ", ("GPX", 'b'), " or ", ("KML", 'b'), " and a file name, then Save."],
    ])
    S.body(doc, ["Your lines and areas are written out in the chosen format — in ", ("GPX", 'b'),
        ", lines are saved as tracks and areas as routes; in ", ("KML", 'b'),
        ", they are saved as line and polygon placemarks."])
    S.callout(doc, "important", "What does not export.",
        "Circles and text labels are not included in exported files — neither format has a clean equivalent, so they are "
        "skipped. If you need a circle to travel to another program, trace it with the polygon tool before exporting.")
    S.callout(doc, "important", "Drawings are not saved automatically.",
        "Your drawings live only for the current session and are cleared when you close APRS Command. To keep a set of "
        "drawings, Export them to a GPX or KML file and Import that file next time.")

    S.h1(doc, "Troubleshooting")
    S.bullets(doc, [
        [("“I can’t pan the map.”", 'b'), " You are in draw mode — dragging draws instead of panning. Click ✕ Exit on the "
         "toolbar (you can still zoom with the wheel while drawing)."],
        [("“My circle didn’t appear.”", 'b'), " Press at the center and drag outward before releasing; a click with no drag makes no circle."],
        [("“My color change affected only one tool.”", 'b'), " That is by design — color, fill, and width are remembered "
         "per tool. Switch to the other tool and set its color there."],
        [("“I lost my drawings after restarting.”", 'b'), " Drawings are session-only. Export them before closing to keep them."],
    ])


def ch_replay(doc, n):
    S.chapter_open(doc, n, "Replay",
        "Play recorded APRS traffic back on the map — safely, at your own pace.",
        ["What Replay is and why to use it", "Recording a log", "Opening & loading",
         "Playing, speed & the control bar", "Returning to live", "Troubleshooting"])

    S.h1(doc, "What Replay Is (and Why You Would Use It)")
    S.body(doc, "Replay lets you take a recording of APRS radio traffic and play it back on the map, exactly as if it "
        "were arriving live — except nothing is transmitted and no real radio is involved. Think of it like a DVR for "
        "your APRS traffic: you record what came in, then watch it again later.")
    S.body(doc, "There are two common reasons to use it:")
    S.bullets(doc, [
        [("Review what you may have missed.", 'b'), " If you stepped away during a busy net or an incident, you can "
         "replay that stretch of time and watch every station, object, and message appear again — at your own pace."],
        [("Test your setup safely.", 'b'), " Because replay runs through the same machinery as live traffic, it is a "
         "safe way to see how your alerts, filters, and the map behave — without keying a radio or waiting for real activity."],
    ])
    S.body(doc, "Replay is one of three related, transmit-safe tools — Replay, Simulation, and Training. This chapter "
        "covers Replay; the others have their own chapters.")
    S.callout(doc, "important", "IMPORTANT — Replay never transmits.",
        "Replay only shows data on your own screen. It never sends anything over the air (RF) or to the internet "
        "(APRS-IS), and it is never mixed into anyone else’s traffic. The Replay window even shows “Replay "
        "transmit disabled” so you always know. You cannot accidentally transmit while replaying.")

    S.h1(doc, "Before You Begin: Recording a Log to Replay")
    S.body(doc, "You cannot replay traffic you never recorded. A replay file is created by the Raw Packet Monitor, which "
        "quietly captures every packet your active connections receive. “Recording” simply means leaving it "
        "running, then saving what it captured to a file.")
    S.body(doc, [("To record a log:", 'b')])
    S.steps(doc, [
        ["Make sure a source is connected and receiving — for example an ", ("APRS-IS", 'b'), " internet connection or an ",
         ("RF/TNC", 'b'), " port (see the Connections chapter). You need real traffic coming in to record it."],
        ["Open ", ("View → Raw Packets", 'b'), " to bring up the Raw Packet Monitor. Every packet that arrives "
         "appears here in the order it was received — that live list is your proof that packets are flowing."],
        ["Leave it running and collecting for as long as you want the recording to cover."],
        ["Click ", ("Save Log", 'b'), " (bottom-left of the Raw Packet Monitor) and choose a file name and location. The "
         "file is saved with an ", (".aprslog", 'b'), " extension. A short message confirms how many packets were saved, "
         "e.g. “Saved 128 packets to …”."],
    ])
    S.screenshot(doc, "Raw Packet Monitor with packets listed and the Save Log button highlighted")
    S.callout(doc, "tip", "TIP",
        "The Raw Packet Monitor is always the “recorder” — there is no separate Record button. Save Log is "
        "simply the “stop and save what I’ve captured” step. If you click Save Log with nothing captured "
        "yet, it will tell you so instead of writing an empty file.")

    S.h1(doc, "Opening the Replay Window")
    S.steps(doc, [
        ["From the ", ("menu bar", 'b'), ", choose ", ("Tools → Replay", 'b'), "."],
        ["The Replay window opens, titled ", ("“Replay Mode”", 'b'), ", and shows “Replay transmit "
         "disabled” as a reminder that nothing will be transmitted."],
    ])
    S.screenshot(doc, "The Replay window (“Replay Mode”) as it first opens, with Browse, Load, and Play")

    S.h1(doc, "Loading a Log File")
    S.steps(doc, [
        ["Click ", ("Browse…", 'b'), " and select the ", (".aprslog", 'b'), " file you saved earlier. (Your saved "
         "logs appear under the “APRS log files” filter.)"],
        ["Click ", ("Load", 'b'), ". The window reads the file and shows how many packets it contains and a preview list "
         "of the entries. The status line reads ", ("Ready", 'b'), "."],
    ])
    S.screenshot(doc, "A log file loaded — the entry list populated and the status reading “Ready”")

    S.h1(doc, "Playing a Log")
    S.body(doc, "When you press Play, two things happen at once, and it is worth knowing what to expect so nothing surprises you:")
    S.steps(doc, [
        ["Optionally set the ", ("Speed", 'b'), " first (see below). Then click ", ("Play", 'b'), "."],
        ["The map ", ("clears", 'b'), " and begins filling in only with the stations from your recording — so you are "
         "looking at the recorded session cleanly, not mixed with current traffic."],
        ["The Replay window ", ("collapses out of the way", 'b'), ", and a compact ",
         ("control bar appears across the top-center of the map", 'b'), " so you can watch the map while you control playback."],
    ])
    S.callout(doc, "note", "What happens to live traffic while you replay?",
        "It is not lost. Your live connections keep receiving and quietly caching everything in the background. It simply "
        "isn’t shown while you’re reviewing the recording. When you return to live (below), the map shows current "
        "traffic plus everything that arrived while you were replaying.")
    S.screenshot(doc, "The map mid-replay: recorded stations filling in, with the control bar at the top")

    S.h1(doc, "The On-Map Replay Control Bar")
    S.body(doc, "The control bar is your “remote” for the recording. Depending on whether playback is running "
        "or finished, you will see:")
    S.bullets(doc, [
        [("Pause / Resume", 'b'), " — temporarily halts playback in place, or continues it. (You see Pause while it is playing, Resume while it is paused.)"],
        [("Stop", 'b'), " — ends the current playback run (the recorded stations stay on the map for you to study)."],
        [("Replay again", 'b'), " — appears once a run finishes or is stopped; clears the map and plays the same log from the start."],
        [("Return to Live", 'b'), " — ends the review and switches the map back to live traffic (see below)."],
        [("Speed", 'b'), " — the playback speed multiplier (see below)."],
        [("Position and progress", 'b'), " — how far through the recording you are (for example “42 / 128” and “47%”)."],
    ])
    S.screenshot(doc, "Close-up of the control bar with each button labeled")

    S.h1(doc, "How Fast Does It Play? (Speed)")
    S.body(doc, "By default, replay runs at real time (1×): if two packets were originally 10 seconds apart, they "
        "appear 10 seconds apart. This faithfully re-creates how the session unfolded — but a long recording takes just "
        "as long to watch.")
    S.body(doc, ["To speed it up, raise the ", ("Speed", 'b'), " number on the control bar. At ", ("10×", 'b'),
        ", a recording that spanned ten minutes plays in about one minute. You can change speed at any time during "
        "playback. (Silent gaps are capped so a long quiet stretch never stalls the replay.)"])
    S.callout(doc, "note", "NOTE",
        "Speed resets to 1× each time you start a new review, so you always begin at real time and choose to speed up from there.")

    S.h1(doc, "Returning to Live")
    S.body(doc, "When you are done reviewing:")
    S.steps(doc, [
        ["Click ", ("Return to Live", 'b'), " on the control bar."],
        ["The map switches back to ", ("live traffic", 'b'), ", showing current stations ",
         ("plus everything that arrived while you were replaying", 'b'), " (it was cached the whole time)."],
        ["The full Replay window reappears, so you can load another log if you wish."],
    ])

    S.h1(doc, "What You Will See on the Map During Replay")
    S.bullets(doc, [
        ["Replayed stations are ", ("tagged as “Replay”", 'b'), " in the Station List and Raw Packet Monitor, "
         "so you can always tell recorded data from live data."],
        ["The Station List follows the map — during replay it shows the ", ("recorded", 'b'), " stations, and it switches "
         "back to live when you return to live."],
        ["A station only appears if the recording contained a ", ("position", 'b'), " for it. A log of 100 packets often "
         "produces fewer than 100 pins, because some packets carry no location and several packets can come from the same station."],
    ])

    S.h1(doc, "Troubleshooting")
    S.bullets(doc, [
        [("“Nothing happened when I pressed Play.”", 'b'), " Make sure a log is loaded (the status reads Ready "
         "and the entry list is populated). An empty or non-position log has nothing to show."],
        [("“It’s playing very slowly.”", 'b'), " That is real-time (1×) pacing. Raise the Speed on the control bar to 10× or 20×."],
        [("“Fewer stations appeared than packets in the log.”", 'b'), " Expected — only packets with a position "
         "become map pins, and repeated packets from one station update a single pin. Open the Station List (filtered to "
         "“Replay”) to see everything the log produced."],
        [("“I closed the Replay window and it won’t come back on Return to Live.”", 'b'), " If you close the "
         "Replay window with its ×, reopen it from the Replay button; it only hides (not closes) automatically during a review."],
    ])


# Hand-written chapters, each tagged with its outline order. JSON chapters (docs/manual-build/chapters/
# *.json, each with an "order") are merged in by order, so chapters slot into the right place regardless
# of whether they're Python functions or agent-produced JSON. Final numbers are assigned 1..N by order.
FUNCTION_CHAPTERS = [
    (1,  ch_welcome),
    (2,  ch_installing),
    (4,  ch_maptour),
    (7,  ch_drawing),
    (25, ch_replay),
]


def load_json_chapters():
    import glob, json
    out = []
    folder = os.path.join(os.path.dirname(__file__), "chapters")
    for path in sorted(glob.glob(os.path.join(folder, "*.json"))):
        with open(path, encoding="utf-8") as f:
            ch = json.load(f)
        out.append((int(ch.get("order", 999)), ch))
    return out


def main():
    doc = S.new_document()
    title_page(doc)
    philosophy(doc)
    how_to_use(doc)
    contents(doc)
    amendments_register(doc)

    items = [(order, ("func", fn)) for order, fn in FUNCTION_CHAPTERS]
    items += [(order, ("json", ch)) for order, ch in load_json_chapters()]
    items.sort(key=lambda x: x[0])

    for number, (_order, (kind, payload)) in enumerate(items, 1):
        if kind == "func":
            payload(doc, number)
        else:
            S.render_chapter(doc, payload, number)

    out = os.environ.get("MANUAL_OUT") or os.path.join(os.path.dirname(__file__), "..", "USER_MANUAL.docx")
    out = os.path.abspath(out)
    doc.save(out)
    print("OK wrote", out, "—", len(items), "chapters")


if __name__ == "__main__":
    main()
