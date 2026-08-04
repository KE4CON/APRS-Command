# -*- coding: utf-8 -*-
"""
Shared styling + building blocks for the APRS Command User Manual.

The manual's *source of truth* is this build pipeline (style.py + build.py + the
chapter modules). Running build.py regenerates docs/USER_MANUAL.docx from scratch.

IMPORTANT WORKFLOW NOTE
-----------------------
Regenerating overwrites the .docx. Screenshots are added by hand in Word only
AFTER every chapter's text is complete — until then, regenerating is safe.
Once screenshots are in, do NOT regenerate; make further changes in Word, and
ship future changes as dated amendment supplements (see the amendment model).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette --------------------------------------------------------------
INK      = RGBColor(0x22, 0x2B, 0x38)
ACCENT   = RGBColor(0x1F, 0x4E, 0x79)   # deep blue
ACCENT2  = RGBColor(0x2E, 0x6D, 0xA4)   # medium blue
MUTED    = RGBColor(0x64, 0x74, 0x8B)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_HEX = "1F4E79"

CALLOUTS = {
    "important": ("FFF4CC", "E0A800", RGBColor(0x8A, 0x6D, 0x00)),
    "note":      ("E7F0FA", "2E6DA4", ACCENT),
    "tip":       ("E8F5E9", "43A047", RGBColor(0x2E, 0x7D, 0x32)),
    "warning":   ("FDECEA", "D9534F", RGBColor(0xA3, 0x2A, 0x2A)),
}
SHOT_FILL, SHOT_BORDER = "F1F5F9", "94A3B8"

# ---- low-level xml helpers ------------------------------------------------
def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _borders(cell, sides):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in sides:
            s = sides[edge]
            el = OxmlElement('w:' + edge)
            el.set(qn('w:val'), s.get('val', 'single'))
            el.set(qn('w:sz'), str(s.get('sz', 6)))
            el.set(qn('w:space'), str(s.get('space', 0)))
            el.set(qn('w:color'), s.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def _margins(cell, top=110, start=170, bottom=110, end=170):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        el = OxmlElement('w:' + m); el.set(qn('w:w'), str(v)); el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def par_border(p, edge, sz=4, color='CBD5E1', val='single', space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    el = OxmlElement('w:' + edge)
    el.set(qn('w:val'), val); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), str(space)); el.set(qn('w:color'), color)
    pbdr.append(el)

def _field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); run._r.append(b)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; run._r.append(it)
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate'); run._r.append(sep)
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); run._r.append(e)

# ---- run helpers ----------------------------------------------------------
def add_runs(p, parts):
    """parts: string, or list of str / (text, 'bi') tuples."""
    if not isinstance(parts, list):
        parts = [parts]
    for part in parts:
        if isinstance(part, tuple):
            t, st = part; r = p.add_run(t)
            if 'b' in st: r.bold = True
            if 'i' in st: r.italic = True
            if 'c' in st:  # code / monospace-ish
                r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        else:
            p.add_run(part)

# ---- document setup -------------------------------------------------------
def new_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
    sec.different_first_page_header_footer = True  # clean title page

    normal = doc.styles['Normal']
    normal.font.name = 'Segoe UI'; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.15; normal.paragraph_format.space_after = Pt(8)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI Semibold'; h1.font.size = Pt(19); h1.font.color.rgb = ACCENT; h1.font.bold = False
    h1.paragraph_format.space_before = Pt(18); h1.paragraph_format.space_after = Pt(6)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI Semibold'; h2.font.size = Pt(13.5); h2.font.color.rgb = ACCENT2; h2.font.bold = False
    h2.paragraph_format.space_before = Pt(14); h2.paragraph_format.space_after = Pt(4)

    # running header (right) + footer page number (center) on non-title pages
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("APRS Command  ·  User Manual"); hr.font.name = 'Segoe UI'; hr.font.size = Pt(8.5); hr.font.color.rgb = MUTED
    par_border(hp, 'bottom', sz=4, color='D7DEE7', space=4)
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_border(fp, 'top', sz=4, color='D7DEE7', space=4)
    fr = fp.add_run("Page "); fr.font.size = Pt(8.5); fr.font.color.rgb = MUTED
    _field(fp, ' PAGE ')
    return doc

# ---- content blocks -------------------------------------------------------
def body(doc, parts):
    p = doc.add_paragraph(); add_runs(p, parts); return p

def h1(doc, text):
    return doc.add_heading(text, level=1)

def h2(doc, text):
    return doc.add_heading(text, level=2)

def step(doc, n, parts):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.38); pf.first_line_indent = Inches(-0.38); pf.space_after = Pt(5)
    pf.tab_stops.add_tab_stop(Inches(0.38), WD_TAB_ALIGNMENT.LEFT)
    r = p.add_run("%d.\t" % n); r.bold = True; r.font.color.rgb = ACCENT
    add_runs(p, parts)

def steps(doc, items):
    for i, it in enumerate(items, 1):
        step(doc, i, it)

def bullet(doc, parts):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.38); pf.first_line_indent = Inches(-0.20); pf.space_after = Pt(4)
    pf.tab_stops.add_tab_stop(Inches(0.38), WD_TAB_ALIGNMENT.LEFT)
    r = p.add_run(u"▪\t"); r.font.color.rgb = ACCENT2
    add_runs(p, parts)

def bullets(doc, items):
    for it in items:
        bullet(doc, it)

def _one_cell(doc, width=Inches(6.5)):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; tbl.allow_autofit = False
    cell = tbl.cell(0, 0); cell.width = width
    return tbl, cell

def callout(doc, kind, label, text):
    fill, border, labelcolor = CALLOUTS[kind]
    _, cell = _one_cell(doc)
    _shade(cell, fill)
    _borders(cell, {'left': {'sz': 30, 'color': border}, 'top': {'sz': 4, 'color': border},
                    'bottom': {'sz': 4, 'color': border}, 'right': {'sz': 4, 'color': border}})
    _margins(cell)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    rl = p.add_run(label + "   "); rl.bold = True; rl.font.color.rgb = labelcolor; rl.font.size = Pt(10.5)
    add_runs(p, text if isinstance(text, list) else [text])
    for r in p.runs[1:]:
        if r.font.color.rgb is None:
            r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def screenshot(doc, caption):
    _, cell = _one_cell(doc)
    _shade(cell, SHOT_FILL)
    dashed = {'val': 'dashed', 'sz': 8, 'color': SHOT_BORDER}
    _borders(cell, {'left': dashed, 'top': dashed, 'bottom': dashed, 'right': dashed})
    _margins(cell, top=220, bottom=220)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SCREENSHOT"); r.bold = True; r.font.color.rgb = MUTED; r.font.size = Pt(9)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(caption); r2.italic = True; r2.font.color.rgb = MUTED; r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def page_break(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

# ---- chapter opener -------------------------------------------------------
def chapter_open(doc, number, title, subtitle, in_this_chapter=None):
    page_break(doc)
    kick = doc.add_paragraph(); kick.paragraph_format.space_after = Pt(2)
    kr = kick.add_run("CHAPTER %d" % number); kr.bold = True; kr.font.size = Pt(10)
    kr.font.color.rgb = ACCENT2; kr.font.name = 'Segoe UI Semibold'
    t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(2)
    tr = t.add_run(title); tr.font.name = 'Segoe UI Semibold'; tr.font.size = Pt(28); tr.font.color.rgb = ACCENT
    # register as Heading 1 outline level so it shows in the TOC
    _set_outline_level(t, 0)
    _style_as_toc_entry(t, title, level=1)
    if subtitle:
        s = doc.add_paragraph()
        sr = s.add_run(subtitle); sr.italic = True; sr.font.size = Pt(11.5); sr.font.color.rgb = MUTED
        par_border(s, 'bottom', sz=12, color=ACCENT_HEX, space=8)
    if in_this_chapter:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
        itc = doc.add_paragraph(); itc.paragraph_format.space_after = Pt(2)
        ir = itc.add_run("IN THIS CHAPTER"); ir.bold = True; ir.font.size = Pt(9)
        ir.font.color.rgb = MUTED; ir.font.name = 'Segoe UI Semibold'
        for item in in_this_chapter:
            b = doc.add_paragraph(); b.paragraph_format.left_indent = Inches(0.2); b.paragraph_format.space_after = Pt(1)
            rb = b.add_run(u"•  " + item); rb.font.size = Pt(10); rb.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def _set_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(level)); pPr.append(ol)

def _style_as_toc_entry(p, text, level=1):
    # Ensure the paragraph carries the built-in Heading style name so TOC \u picks it up,
    # while keeping our custom visual formatting via direct run formatting above.
    pPr = p._p.get_or_add_pPr()
    ps = pPr.find(qn('w:pStyle'))
    if ps is None:
        ps = OxmlElement('w:pStyle'); pPr.insert(0, ps)
    ps.set(qn('w:val'), 'Heading1')

def toc(doc):
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-2" \\h \\z \\u')

def _table_light_borders(t, color="D7DEE7", sz=4):
    tblPr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

def table(doc, headers, rows, widths):
    """Styled table: accent header row, zebra body rows, light borders.
    Each cell value is an add_runs-style string or list of runs."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.allow_autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = widths[i]
        _shade(c, ACCENT_HEX)
        _margins(c, top=70, bottom=70, start=120, end=120)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h); r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10); r.font.name = 'Segoe UI Semibold'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]; c.width = widths[i]
            if ri % 2 == 1: _shade(c, "F3F6FA")
            _margins(c, top=60, bottom=60, start=120, end=120)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            add_runs(p, val)
            for rr in p.runs:
                if rr.font.size is None: rr.font.size = Pt(9.5)
    _table_light_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t
