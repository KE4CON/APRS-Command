"""
APRS Command — one-page Fact Sheet builder (navy+gold, matching the manual/guide/quick-start).

A single professionally styled page selling APRS Command to the ham radio community. Run:
  python factsheet_build.py    ->  writes ../APRS_Command_Fact_Sheet.docx
Env: FACTSHEET_OUT overrides the output path.
"""
import os, sys, datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT as VAL

HERE = os.path.dirname(__file__)
sys.path.insert(0, os.path.join(HERE, "..", "manual-build"))
import style as S  # noqa: E402

CONTENT_W = 7.3  # inches, with 0.6" side margins on US Letter


def band(doc, height_in, fill):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; tbl.allow_autofit = False
    cell = tbl.cell(0, 0); cell.width = Inches(CONTENT_W)
    S._shade(cell, fill); S._no_borders(cell)
    S._set_row_height(tbl.rows[0], height_in)
    cell.vertical_alignment = VAL.CENTER
    S._margins(cell, top=90, bottom=90, start=240, end=240)
    return cell


def centered(cell, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    return p


def run(p, text, *, size, color, bold=False, name="Segoe UI"):
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.font.name = name
    return r


def feature_cell(cell, title, bullets):
    S._shade(cell, S.PANEL_HEX); S._no_borders(cell)
    S._margins(cell, top=90, bottom=90, start=150, end=150)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    run(p, title, size=10.5, color=S.ACCENT, bold=True, name="Segoe UI Semibold")
    for b in bullets:
        bp = cell.add_paragraph(); bp.paragraph_format.space_after = Pt(1)
        run(bp, "▪  ", size=9, color=S.GOLD)
        run(bp, b, size=9, color=S.INK)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.5)
    sec.left_margin = sec.right_margin = Inches(0.6)
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"; normal.font.size = Pt(10); normal.font.color.rgb = S.INK
    normal.paragraph_format.line_spacing = 1.06; normal.paragraph_format.space_after = Pt(3)

    # ── Gold top rule ─────────────────────────────────────────────────────────
    top = doc.add_paragraph(); top.paragraph_format.space_after = Pt(0)
    S.par_border(top, "bottom", sz=22, color=S.GOLD_HEX, space=1)

    # ── Navy header band ──────────────────────────────────────────────────────
    h = band(doc, 1.55, S.NAVY_HEX)
    p = centered(h, first=True); p.paragraph_format.space_after = Pt(2)
    run(p, "APRS COMMAND", size=32, color=S.WHITE, bold=True, name="Segoe UI Semibold")
    p = centered(h); p.paragraph_format.space_after = Pt(4)
    run(p, "Situational Awareness for Amateur Radio", size=13.5, color=S.GOLD, bold=True, name="Segoe UI Semibold")
    p = centered(h)
    run(p, "Free   ·   Open-Source (GPL v3)   ·   Windows · macOS · Linux · Raspberry Pi",
        size=9.5, color=S.COVER_MUTE)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Value proposition ─────────────────────────────────────────────────────
    vp = doc.add_paragraph(); vp.alignment = WD_ALIGN_PARAGRAPH.CENTER; vp.paragraph_format.space_after = Pt(2)
    run(vp, "Turn scattered radio traffic into one live picture of who is where, what the weather is doing, "
            "and where the resources and hazards are.", size=11.5, color=S.ACCENT, bold=True, name="Segoe UI Semibold")
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after = Pt(8)
    run(sub, "A modern, cross-platform APRS client for everyday operating and serious emergency "
             "communications — the open-source successor to UI-View32, the classic APRS client that "
             "went dark in 2004.",
        size=10, color=S.MUTED)

    # ── Feature grid (2 x 2 cards) ────────────────────────────────────────────
    grid = doc.add_table(rows=2, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False; grid.allow_autofit = False
    for r in grid.rows:
        S._set_row_height(r, 0.02, rule="atLeast")
        for c in r.cells:
            c.width = Inches(CONTENT_W / 2)
    feature_cell(grid.cell(0, 0), "SEE EVERYTHING", [
        "Live map of every station, object, and weather report you hear",
        "Sortable station list, raw-packet monitor, movement trails, range rings",
        "Street, topo, and aerial base maps — plus a national weather-radar overlay",
    ])
    feature_cell(grid.cell(0, 1), "COMMUNICATE & COORDINATE", [
        "APRS messaging with templates, delivery tracking, and scheduling",
        "Objects, items, and bulletins; tactical callsigns",
        "Net Control roster, geofence alerts, and after-action ICS-213/214/309 export",
    ])
    feature_cell(grid.cell(1, 0), "OPERATE ANYWHERE", [
        "Internet (APRS-IS) and radio (KISS, Direwolf, AGWPE) side by side",
        "GPS location, downloadable offline maps, iGate and digipeater",
        "Runs identically on a desktop, a laptop, or a Raspberry Pi in the field",
    ])
    feature_cell(grid.cell(1, 1), "PLAN, PRACTICE & EXTEND", [
        "Coverage prediction, elevation profiles, frequency reference, map drawing",
        "Replay, Simulation, and Training — practice with nothing on the air",
        "Open platform: local REST API, WebSocket streams, and plugins",
    ])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Why it stands out ─────────────────────────────────────────────────────
    wh = doc.add_paragraph(); wh.paragraph_format.space_after = Pt(3)
    run(wh, "WHY APRS COMMAND", size=11, color=S.ACCENT, bold=True, name="Segoe UI Semibold")
    S.par_border(wh, "bottom", sz=10, color=S.GOLD_HEX, space=3)
    for lead, rest in [
        ("Receive-first and transmit-safe.",
         " It listens by default and puts nothing on the air until you deliberately enable it. Every "
         "transmission passes through one safety gate — with exercise modes and one-switch EXERCISE "
         "traffic marking for drills."),
        ("Built for EmComm, not bolted on.",
         " Net control, geofencing, after-action reports, weather-station ingest, and training tools that "
         "very few clients ship — in one polished, modern interface."),
        ("Free, and free to stay alive.",
         " Released under the GPL so it can never be taken closed or orphaned. If the author steps away "
         "tomorrow, any operator in the world can carry it forward."),
    ]:
        b = doc.add_paragraph(); b.paragraph_format.space_after = Pt(2); b.paragraph_format.left_indent = Inches(0.02)
        run(b, "✔  ", size=10, color=S.GOLD, bold=True)
        run(b, lead, size=10, color=S.ACCENT, bold=True)
        run(b, rest, size=10, color=S.INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Call-to-action band ───────────────────────────────────────────────────
    cta = band(doc, 0.75, S.NAVY_HEX)
    p = cta.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
    run(p, "Download free at  github.com/KE4CON/APRS-Command", size=12.5, color=S.WHITE, bold=True, name="Segoe UI Semibold")
    p2 = cta.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(0)
    run(p2, "Open-source (GPL v3)  ·  by James Rospopo, KE4CON  ·  73", size=9.5, color=S.GOLD)

    out = os.environ.get("FACTSHEET_OUT") or os.path.join(HERE, "..", "APRS_Command_Fact_Sheet.docx")
    out = os.path.abspath(out)
    doc.save(out)
    print("OK — wrote", out)


if __name__ == "__main__":
    build()
